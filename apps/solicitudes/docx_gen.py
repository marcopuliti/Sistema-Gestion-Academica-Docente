import os
from io import BytesIO

from django.conf import settings
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ESTADO_LABEL = {
    'pendiente':   'en trámite de aprobación',
    'en_revision': 'en revisión',
    'aprobado':    'aprobado',
    'rechazado':   'rechazado',
}


# ── Helpers de estilo ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _p_cell(cell, texto, bold=False, center=False, size=9, underline=False):
    """Escribe texto en una celda con formato."""
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(texto or '')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    _cell_margins(cell)


def _merge_row(table, row_idx, col_start, col_end):
    """Fusiona celdas horizontalmente."""
    row = table.rows[row_idx]
    row.cells[col_start].merge(row.cells[col_end])


def _set_col_widths(table, widths_cm):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])


def _header_row(table, row_idx, labels, bold=True, center=True, bg='D9D9D9'):
    row = table.rows[row_idx]
    for i, label in enumerate(labels):
        _p_cell(row.cells[i], label, bold=bold, center=center)
        if bg:
            _set_cell_bg(row.cells[i], bg)


def _sec_titulo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.underline = True


def _caja_texto(doc, texto):
    """Párrafo con borde simulado mediante tabla de una celda."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    cell.text = ''
    para = cell.paragraphs[0]
    run = para.add_run(texto or '')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    _cell_margins(cell, top=80, bottom=80, left=100, right=100)


def _encabezado_parrafo(doc, texto, bold=False, size=11, center=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold


# ── Generador principal ────────────────────────────────────────────────────────

def generar_docx_solicitud(solicitud):
    doc = Document()

    # Márgenes de página
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── ENCABEZADO ─────────────────────────────────────────────────────────────
    # Tabla 2 columnas: institución izquierda, estado derecha
    t_cab = doc.add_table(rows=1, cols=2)
    t_cab.style = 'Table Grid'

    # Celda izquierda: logo + líneas institucionales
    cell_left = t_cab.cell(0, 0)
    cell_left.text = ''
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'escudo.gif')
    if os.path.exists(logo_path):
        p_logo = cell_left.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(logo_path, width=Cm(2.8))

    departamento = (
        (solicitud.usuario.departamento if solicitud.usuario else None)
        or solicitud.departamento_docente
        or '—'
    )
    for linea, bold in [
        ('Ministerio de Cultura y Educación', True),
        ('Universidad Nacional de San Luis', True),
        ('Facultad de Ciencias Físico Matemáticas y Naturales', True),
        (f'Departamento: {departamento}', True),
    ]:
        p = cell_left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(linea)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = bold

    if solicitud.area:
        p = cell_left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f'Area: {solicitud.area}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True

    # Celda derecha: estado
    year = solicitud.fecha_inicio.year
    estado_txt = ESTADO_LABEL.get(solicitud.estado, solicitud.estado)
    fecha_pres = solicitud.fecha_creacion.strftime('%d/%m/%Y %H:%M:%S')

    cell_right = t_cab.cell(0, 1)
    cell_right.text = ''
    # Spacer vacío para alinear verticalmente con logo
    p_sp = cell_right.paragraphs[0]
    p_sp.paragraph_format.space_before = Pt(30)
    for linea in [
        f'(Programa del año {year})',
        f'(Programa {estado_txt})',
        f'(Presentado el {fecha_pres})',
    ]:
        p = cell_right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(linea)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    # Quitar bordes de la tabla de encabezado
    for row in t_cab.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    _set_col_widths(t_cab, [10.6, 6.4])
    doc.add_paragraph()

    # ── I — Oferta Académica ───────────────────────────────────────────────────
    _sec_titulo(doc, 'I - Oferta Académica')

    t1 = doc.add_table(rows=2, cols=5)
    t1.style = 'Table Grid'
    _header_row(t1, 0, ['Materia', 'Carrera', 'Plan', 'Año', 'Período'])
    row1 = t1.rows[1]
    _p_cell(row1.cells[0], solicitud.nombre_curso)
    _p_cell(row1.cells[1], solicitud.carrera.nombre if solicitud.carrera else '')
    _p_cell(row1.cells[2], solicitud.plan_estudio.codigo if solicitud.plan_estudio else '')
    _p_cell(row1.cells[3], solicitud.anno_carrera or '')
    _p_cell(row1.cells[4], solicitud.get_periodo_display() if solicitud.periodo else '')
    _set_col_widths(t1, [4.5, 4.5, 2.2, 1.6, 3.2])

    doc.add_paragraph()

    # ── II — Equipo Docente ────────────────────────────────────────────────────
    _sec_titulo(doc, 'II - Equipo Docente')

    miembros = list(solicitud.equipo_docente.all())
    n_filas = max(len(miembros), 1) + 1
    t2 = doc.add_table(rows=n_filas, cols=4)
    t2.style = 'Table Grid'
    _header_row(t2, 0, ['Docente', 'Función', 'Cargo', 'Dedicación'])
    for i, m in enumerate(miembros):
        row = t2.rows[i + 1]
        _p_cell(row.cells[0], m.nombre)
        _p_cell(row.cells[1], m.get_funcion_display())
        _p_cell(row.cells[2], m.get_cargo_display() if m.cargo else '')
        _p_cell(row.cells[3], m.get_dedicacion_display() if m.dedicacion else '')
    if not miembros:
        for c in t2.rows[1].cells:
            _p_cell(c, '')
    _set_col_widths(t2, [6.4, 4.5, 2.6, 2.5])

    doc.add_paragraph()

    # ── III — Características del Curso ───────────────────────────────────────
    _sec_titulo(doc, 'III - Características del Curso')

    total_hs = solicitud.total_horas_semanales
    cant_hs  = solicitud.cantidad_semanas * total_hs

    # Tabla 1: Crédito Horario Semanal
    t3a = doc.add_table(rows=3, cols=5)
    t3a.style = 'Table Grid'
    _merge_row(t3a, 0, 0, 4)
    _p_cell(t3a.rows[0].cells[0], 'Crédito Horario Semanal', bold=True, center=True)
    _set_cell_bg(t3a.rows[0].cells[0], 'D9D9D9')
    _header_row(t3a, 1, [
        'Teórico/Práctico', 'Teóricas', 'Prácticas de Aula',
        'Práct. de lab/ camp/ Resid/ PIP, etc.', 'Total',
    ])
    row = t3a.rows[2]
    for cell, val in zip(row.cells, [
        'Hs',
        f'{solicitud.hs_teoricas} Hs',
        f'{solicitud.hs_practicas_aula} Hs',
        f'{solicitud.hs_lab_campo} Hs',
        f'{total_hs} Hs',
    ]):
        _p_cell(cell, val, center=True)
    _set_col_widths(t3a, [2.9, 2.2, 3.2, 5.5, 2.2])

    doc.add_paragraph()

    # Tabla 2: Tipificación / Período
    t3b = doc.add_table(rows=2, cols=2)
    t3b.style = 'Table Grid'
    _header_row(t3b, 0, ['Tipificación', 'Período'])
    row = t3b.rows[1]
    _p_cell(row.cells[0], solicitud.get_modalidad_cursado_display() if solicitud.modalidad_cursado else '')
    _p_cell(row.cells[1], solicitud.get_periodo_display() if solicitud.periodo else '')
    _set_col_widths(t3b, [8.0, 8.0])

    doc.add_paragraph()

    # Tabla 3: Duración
    t3c = doc.add_table(rows=3, cols=4)
    t3c.style = 'Table Grid'
    _merge_row(t3c, 0, 0, 3)
    _p_cell(t3c.rows[0].cells[0], 'Duración', bold=True, center=True)
    _set_cell_bg(t3c.rows[0].cells[0], 'D9D9D9')
    _header_row(t3c, 1, ['Desde', 'Hasta', 'Cantidad de Semanas', 'Cantidad de Horas'])
    row = t3c.rows[2]
    _p_cell(row.cells[0], solicitud.fecha_inicio.strftime('%d/%m/%Y'), center=True)
    _p_cell(row.cells[1], solicitud.fecha_hasta.strftime('%d/%m/%Y'), center=True)
    _p_cell(row.cells[2], str(solicitud.cantidad_semanas), center=True)
    _p_cell(row.cells[3], str(cant_hs), center=True)
    _set_col_widths(t3c, [4.0, 4.0, 4.0, 4.0])

    # ── IV–XIV — Secciones de texto ────────────────────────────────────────────
    secciones = [
        ('IV',    'Fundamentación',                  solicitud.fundamentacion),
        ('V',     'Objetivos / Resultados de Aprendizaje', solicitud.objetivos),
        ('VI',    'Contenidos',
            ('Contenidos Mínimos\n' + solicitud.contenidos_minimos + '\n\n' if solicitud.contenidos_minimos else '')
            + solicitud.unidades),
        ('VII',   'Plan de Trabajos Prácticos',      solicitud.plan_trabajos_practicos or ''),
        ('VIII',  'Régimen de Aprobación',            solicitud.regimen_aprobacion),
        ('IX',    'Bibliografía Básica',              solicitud.bibliografia_basica),
        ('X',     'Bibliografía Complementaria',      solicitud.bibliografia_complementaria or ''),
        ('XI',    'Resumen de Objetivos',             solicitud.resumen_objetivos or ''),
        ('XII',   'Resumen del Programa',             solicitud.resumen_programa or ''),
        ('XIII',  'Imprevistos',                      solicitud.imprevistos or ''),
        ('XIV',   'Otros',                            solicitud.contacto_otros or ''),
    ]
    for num, titulo, texto in secciones:
        _sec_titulo(doc, f'{num} - {titulo}')
        _caja_texto(doc, texto)

    # ── ELEVACIÓN y APROBACIÓN ─────────────────────────────────────────────────
    doc.add_page_break()

    t_firma = doc.add_table(rows=5, cols=2)
    t_firma.style = 'Table Grid'
    _merge_row(t_firma, 0, 0, 1)
    _p_cell(t_firma.rows[0].cells[0], 'ELEVACIÓN y APROBACIÓN DE ESTE PROGRAMA', bold=True, center=True)
    _set_cell_bg(t_firma.rows[0].cells[0], 'D9D9D9')
    _p_cell(t_firma.rows[1].cells[1], 'Profesor Responsable', bold=True, center=True)
    for i, label in enumerate(['Firma:', 'Aclaración:', 'Fecha:'], start=2):
        _p_cell(t_firma.rows[i].cells[0], label)
        cell = t_firma.rows[i].cells[1]
        cell.text = ''
        tc = cell._tc
        trPr = tc.getparent().get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '500')
        trPr.append(trHeight)
    _set_col_widths(t_firma, [4.8, 11.2])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
